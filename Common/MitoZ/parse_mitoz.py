import re
from Bio import SeqIO
from Bio.Seq import Seq
from Bio.SeqRecord import SeqRecord

from docx import Document
from pathlib import Path
import shutil

def GetAllFilePaths(pwd,wildcard='*'):
    '''
    获取目录下文件全路径，通配符检索特定文件名，返回列表
    param: str  "pwd"
    return:dirname pathlab_obj
    return:list [ str ]
    #https://zhuanlan.zhihu.com/p/36711862
    #https://www.cnblogs.com/sigai/p/8074329.html
    '''
    files_lst = []
    target_path=Path(pwd)
    for child in target_path.rglob(wildcard):
        if child.is_symlink():
            pass
        elif child.is_dir():
            pass
        elif child.is_file():
            files_lst.append(child)
    return files_lst


def get_d_loop_seq(seq_dic,d_loop_pos):
    # seq_dic[] = ""
    d_loop_seq = ""
    lst_len = len(d_loop_pos)
    if lst_len == 1:
        chrom,start,end = d_loop_pos[0]
        d_loop_seq = seq_dic[chrom][start-1:end-1]
        # print("##########TEST#########")
        # print(seq_dic[chrom])
        # print(start,end)
        # print("##########TEST#########")
    elif lst_len == 2:
        # first_index end_index # 前面控制了chrom1 chrom2 必须一样
        chrom1,start1,end1 = d_loop_pos[0] # first_index
        chrom2,start2,end2 = d_loop_pos[1] # end_index
        end_seq = seq_dic[chrom1][start1:end1-1]
        start_seq = seq_dic[chrom1][start2-1:end2]
        d_loop_seq = start_seq + "N" + end_seq
    else: # 等于0 或者大于2
        print("d_loop_pos len illegal",d_loop_pos,len(d_loop_pos))
        pass
    return d_loop_seq

def parse_d_loop_seq(fasta_path,summary_path):
    #seq_obj = SeqIO.parse("mitoz_YZ02.complete.fasta_mitoscaf.fa.gbf.fasta", "fasta")
    seq_obj = SeqIO.parse(str(fasta_path), "fasta")
    seq_dic = {}
    for record in seq_obj:
        seq_id = re.split(";",record.id)[0]
        print(seq_id)
        seq_dic[seq_id]=str(record.seq)
    
    # 1-base
    with open(summary_path,mode='rt') as fh:
        fh.readline()
        info_str = fh.readline()
        inf_record = re.split(r"\s+",info_str)
        sample_name = inf_record[0]
        all_seq_len = int(inf_record[1])
        print(all_seq_len)
        all_string = fh.read()
        # print(all_string)
        # ('KKS01', '86', '383', '298', '+', 'CDS', 'ND4L', 'NADH dehydrogenase subunit 4L', '1')
        all_lst = re.findall(r"("+sample_name+r")\s{2,}(.+?)\s{2,}(.+?)\s{1,}(.+?)\s{2,}(.+?)\s{2,}(.+?)\s{2,}(.+?)\s{2,}(.+?)\s{2,}(.+?)\s{2,}\s+?",all_string)
        #all_lst = re.findall(r"(.+?)\s{2,}(.+?)\s{2,}(.+?)\s{1,}(.+?)\s{2,}(.+?)\s{2,}(.+?)\s{2,}(.+?)\s{2,}(.+?)\s{2,}(.+?)\s{2,}\s+?",all_string)
        # print(all_lst[1])
        # print(all_lst[2])
        # print(all_lst[4])

        all_list=[]
        index_key_lst=[]
        for item in all_lst:
            ##Seq_id        Start  End    Length(bp) Direction  Type   Gene_name  Gene_prodcut   Total_freq_occurred
            seq_id ,start,end,seq_len,direction,seq_type,gene_name,gene_prodcut,total_freq_occurred = item
            #print("#",end)
            #print(seq_id ,start,end,seq_len,direction,seq_type,gene_name,gene_prodcut,total_freq_occurred)
            if "<" in start:
                start = re.sub("<","",start)
            if ">" in end:
                end = re.sub(">","",end)
            if start == "Start":
                continue
            # tRNA-Pro  tRNA-Phe
            if gene_prodcut == "tRNA-Pro" or gene_prodcut == "tRNA-Phe":
                all_list.append([seq_id,int(start),int(end)])
                index_key_lst.append([seq_id,int(start),int(end)])
            else:
                all_list.append([seq_id,int(start),int(end)])
            #print([seq_id,int(start),int(end)])
        list_len = len(all_list)    
        break_index = [0,list_len-1] # 断点在D-loop

        d_loop_pos = []  # 一个元素就是连续坐标，两个元素就是前后片段 [chrom start,end]
        # 获取 目标 index
        result_index_lst = []
        for  key in index_key_lst:
            result_index_lst.append(all_list.index(key))
        if len(index_key_lst)>2:
            print("tRNA-Pro 与 tRNA-Phe 不相邻")
        elif len(index_key_lst)==0:
            d_loop_seq=""
            return sample_name,d_loop_seq,d_loop_pos
        # 获取 通过index 的情况,区分获取提取的序列坐标
        #[9, 10]
        #[0, 36]
        first_index = result_index_lst[0]
        end_index = result_index_lst[1]
        # print("result_index_lst",result_index_lst)
        if  result_index_lst == break_index: # 断点在D-loop
            print("break in D-loop")
            chrom1,end1,_ = all_list[first_index]
            chrom2,_,start2 = all_list[end_index]
            if chrom1 == chrom1:
                d_loop_pos.append([chrom1,0,end1]) # first_index
                d_loop_pos.append([chrom2,start2,all_seq_len]) # end_index
                print(chrom1,0,end1)
                print(chrom2,start2,all_seq_len)
            else:
                print("chrom diff",chrom1,chrom2)
        elif result_index_lst[1] - result_index_lst[0] == 1: # 顺序关系
            #chrom,start,end
            chrom1,_,start = all_list[first_index]
            chrom2,end,_ = all_list[end_index]
            if chrom1 == chrom1:
                d_loop_pos.append([chrom1,start,end])
            else:
                print("chrom diff",chrom1,chrom2)
        else:  # 有问题的,输出情况
            print("index ???:",result_index_lst)
    print(d_loop_pos)# [左开右闭)
    
    d_loop_seq =get_d_loop_seq(seq_dic,d_loop_pos)
    return sample_name,d_loop_seq,d_loop_pos

def main():
    current_dir = Path.cwd()
    docx_file = current_dir.joinpath("all_D-loop_seq.docx")
    # fasta_dir = current_dir.joinpath("fasta")
    # fasta_dir.mkdir(parents=True,exist_ok=True)
    # docx_file = current_dir.joinpath("all_D-loop_seq.docx")
    sample_lst = []
    d_loop_pos_dic = {}
    d_loop_seq_dic = {}
    # >KKS04;s-rRNA;len=948;[8598:9546](+)
    file_paths = GetAllFilePaths(current_dir,wildcard='mitoz_*.gbf')
    for gbff_file in file_paths:
        mitochondrion_fasta_path = gbff_file.with_suffix(".gbf.fasta")
        
        dir_path = gbff_file.parent
        summary_path = dir_path.joinpath("summary.txt")
        D_loop_path = gbff_file.with_suffix(".gbf.D-Loop.fasta")

        sample_name,D_loop_seq,d_loop_pos = parse_d_loop_seq(mitochondrion_fasta_path,summary_path)
        if len(D_loop_seq)!=0:
            d_loop_pos_dic[sample_name] = d_loop_pos
            d_loop_seq_dic[sample_name] = D_loop_seq
            sample_lst.append(sample_name)
            description_str = ";".join(["len="+str(len(D_loop_seq)),str(d_loop_pos)])
            record = SeqRecord(Seq(D_loop_seq), id=sample_name+".D_loop", description=description_str)
            SeqIO.write(record, D_loop_path, "fasta")
        new_gbff_file = gbff_file.with_suffix(".gbff")
        shutil.move(gbff_file, new_gbff_file)


    word_string = ""
    table_lst = []
    table_lst.append("sample_name\tchrom\tstart\tend")
    for sample_name in sample_lst:
        D_loop_seq = d_loop_seq_dic[sample_name]
        word_string += ">{}.D_loop\n{}\n".format(sample_name,D_loop_seq)

        d_loop_pos = d_loop_pos_dic[sample_name]
        len_pos = len(d_loop_pos)
        if len_pos == 1 :
            chrom,start,end= d_loop_pos[0]
            table_lst.append("{}\t{}\t{}\t{}".format(sample_name,chrom,start,int(end)-1))
        elif len_pos == 2 :
            chrom,start1,end1= d_loop_pos[0]
            _,start2,end2= d_loop_pos[1]
            table_lst.append("{}\t{}\t{}\t{}".format(sample_name,chrom,int(start1)+1,int(end1)-1))
            table_lst.append("{}\t{}\t{}\t{}".format(sample_name,chrom,start2,end2))
        else:
            print(sample_name,"?????","d_loop_pos",d_loop_pos)
            pass
    xls_file = current_dir.joinpath("D-loop_position.xls")
    with open(xls_file,mode='wt',encoding="utf-8") as out:
        for item in table_lst:
            out.write(item+"\n")

    # 保存docx
    doc = Document()
    paragraph = doc.add_paragraph()
    paragraph.add_run(word_string)
    doc.save(docx_file)

if __name__ == '__main__':
    main()


'''
#ReadMe
mitoz_*.fasta_mitoscaf.fa.most_related_species.txt  MitoZ 注释认为最接近的物种 Trachypithecus francoisi
mitoz_*.fasta_mitoscaf.fa.tbl                       NCBI上传所需的tbl格式,与gbff 内容相关
errorsummary.val                                    MitoZ 错误日志,不影响结果可忽视
mitoz_*.fasta_mitoscaf.fa.gbff                      gbff 格式,同gb格式,就是genbank格式的文件;可用SnapGene 打开或作为PhyloSuite输入文件
mitoz_*.fasta_mitoscaf.fa.gbf.cds.fasta             gbff 文件中提取的 cds 序列
mitoz_*.fasta_mitoscaf.fa.gbf.cds_translation.fasta gbff 文件中提取的 cds_translation 序列
mitoz_*.fasta_mitoscaf.fa.gbf.fasta                 gbff 文件中提取的完整线粒体(mitochondrion)序列
mitoz_*.fasta_mitoscaf.fa.gbf.gene.fasta            gbff 文件中提取的 gene 序列
mitoz_*.fasta_mitoscaf.fa.gbf.rrna.fasta            gbff 文件中提取的 rrna 序列
mitoz_*.fasta_mitoscaf.fa.gbf.trna.fasta            gbff 文件中提取的 trna 序列
summary.txt                                         gbff 文件中注释内容的汇总,也包含坐标信息,类似gff注释文件
circos.png                                          gbff 文件中特征的方向绘图文件(png)
circos.svg                                          gbff 文件中特征的方向绘图文件(svg)
mitoz_*.fasta_mitoscaf.fa.gbf.D-Loop.fasta          根据tRNA-Pro 与 tRNA-Phe之间为 D-loop 序列特征提取的 D-loop 序列

'''

